"""
report_builder.py
Converts Claude's structured JSON report into a branded Caava Group DOCX.
"""
import io
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

# ---------------------------------------------------------------------------
# Caava Group brand colours
# ---------------------------------------------------------------------------
CAAVA_NAVY   = RGBColor(0x1A, 0x2A, 0x5E)   # dark navy blue
CAAVA_RED    = RGBColor(0xCC, 0x1F, 0x1F)   # Caava red accent
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY   = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY     = RGBColor(0xCC, 0xCC, 0xCC)
DARK_GREY    = RGBColor(0x33, 0x33, 0x33)

SEV_COLOURS = {
    'critical': RGBColor(0xC0, 0x00, 0x00),
    'high':     RGBColor(0xFF, 0x66, 0x00),
    'medium':   RGBColor(0xFF, 0xC0, 0x00),
    'low':      RGBColor(0x00, 0x70, 0xC0),
    'info':     RGBColor(0x70, 0xAD, 0x47),
}

SEV_ORDER = ['critical', 'high', 'medium', 'low', 'info']

LOGO_PATH = Path(__file__).parent / 'images' / 'caava_logo.png'


# ---------------------------------------------------------------------------
# Low-level DOCX helpers
# ---------------------------------------------------------------------------

def _rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(rgb))
    tcPr.append(shd)


def _set_cell_borders(cell, sides=('top', 'bottom', 'left', 'right'), color='CCCCCC', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_space(doc, before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def _heading(doc, text, level=1, colour=None):
    p = doc.add_heading(text, level=level)
    if colour:
        for run in p.runs:
            run.font.color.rgb = colour
    # Level 1 and 2 headings get a red bottom border (matches Caava template style)
    if level in (1, 2):
        _add_bottom_border(p, 'CC1F1F', '8' if level == 1 else '4')
    return p


def _add_run(para, text, bold=False, italic=False, size=None, colour=None, font='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    return run


def _page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Severity bar chart (matplotlib → PNG in memory)
# ---------------------------------------------------------------------------

def _build_severity_chart(findings: list) -> bytes:
    counts = {s: 0 for s in SEV_ORDER}
    for f in findings:
        sev = f.get('severity', 'info').lower()
        if sev in counts:
            counts[sev] += 1

    labels = [s.capitalize() for s in SEV_ORDER]
    values = [counts[s] for s in SEV_ORDER]
    colours = ['#C00000', '#FF6600', '#FFC000', '#0070C0', '#70AD47']

    fig, ax = plt.subplots(figsize=(7, 3))
    bars = ax.bar(labels, values, color=colours, width=0.5, edgecolor='none')
    ax.set_ylabel('Count', fontsize=9)
    ax.set_ylim(0, max(values + [2]) + 1)
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(labelsize=9)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Helper: add a coloured paragraph border (bottom rule)
# ---------------------------------------------------------------------------

def _add_bottom_border(para, colour_hex='CC1F1F', size='12'):
    """Add a bottom border (horizontal rule) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), colour_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_top_border(para, colour_hex='CC1F1F', size='12'):
    """Add a top border (horizontal rule) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), size)
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), colour_hex)
    pBdr.append(top)
    pPr.append(pBdr)


def _set_para_bg(para, colour_hex):
    """Set paragraph background shading."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), colour_hex)
    pPr.append(shd)


# ---------------------------------------------------------------------------
# Cover page — matches Caava Group branding (dark navy background)
# ---------------------------------------------------------------------------

def _build_cover(doc, report_data, target_url, client_name, report_version='1.0'):
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    NAVY_HEX = '1A2A5E'
    RED_HEX  = 'CC1F1F'

    # ── "CONFIDENTIAL" box at top centre (red border, navy bg) ──
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_bg(p_conf, NAVY_HEX)
    pPr = p_conf._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom', 'left', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:color'), RED_HEX)
        pBdr.append(el)
    pPr.append(pBdr)
    run_c = p_conf.add_run('C O N F I D E N T I A L')
    run_c.font.color.rgb = CAAVA_RED
    run_c.font.size      = Pt(8)
    run_c.font.bold      = True
    run_c.font.name      = 'Calibri'

    # ── Navy background spacer ──
    for _ in range(3):
        p_sp = doc.add_paragraph()
        _set_para_bg(p_sp, NAVY_HEX)

    # ── "C AAVA GROUP" heading — red C, white rest ──
    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_bg(p_brand, NAVY_HEX)
    r_c = p_brand.add_run('C')
    r_c.font.color.rgb = CAAVA_RED
    r_c.font.size      = Pt(36)
    r_c.font.bold      = True
    r_c.font.name      = 'Calibri'
    r_rest = p_brand.add_run('AAVA  GROUP')
    r_rest.font.color.rgb = WHITE
    r_rest.font.size      = Pt(36)
    r_rest.font.bold      = True
    r_rest.font.name      = 'Calibri'

    # ── Spacers ──
    for _ in range(2):
        p_sp = doc.add_paragraph()
        _set_para_bg(p_sp, NAVY_HEX)

    # ── Document title (bold white) ──
    exec_sum = report_data.get('executive_summary', {})
    details  = exec_sum.get('assessment_details', {})
    target_label = details.get('target', target_url)
    doc_title = f"VULNERABILITY ASSESSMENT REPORT"
    doc_sub   = target_label.upper()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_bg(p_title, NAVY_HEX)
    r_title = p_title.add_run(doc_title)
    r_title.font.color.rgb = WHITE
    r_title.font.size      = Pt(22)
    r_title.font.bold      = True
    r_title.font.name      = 'Calibri'

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_bg(p_sub, NAVY_HEX)
    r_sub = p_sub.add_run(doc_sub)
    r_sub.font.color.rgb = WHITE
    r_sub.font.size      = Pt(16)
    r_sub.font.bold      = True
    r_sub.font.name      = 'Calibri'

    # ── Red rule ──
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_bg(p_rule, NAVY_HEX)
    _add_bottom_border(p_rule, RED_HEX, '16')

    # ── Subtitle / description ──
    assessment_type = details.get('assessment_type', 'Web Application Security Assessment')
    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_bg(p_desc, NAVY_HEX)
    r_desc = p_desc.add_run(assessment_type)
    r_desc.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r_desc.font.size      = Pt(11)
    r_desc.font.name      = 'Calibri'

    # ── Red rule ──
    p_rule2 = doc.add_paragraph()
    _set_para_bg(p_rule2, NAVY_HEX)
    _add_bottom_border(p_rule2, RED_HEX, '6')

    # ── Spacers ──
    for _ in range(2):
        p_sp = doc.add_paragraph()
        _set_para_bg(p_sp, NAVY_HEX)

    # ── Metadata table (on white background now) ──
    meta_rows = [
        ('Document ID',      f"CAAVA-VAPT-{datetime.now().strftime('%Y%m%d')}"),
        ('Version',          report_version),
        ('Assessment Date',  details.get('assessment_date', datetime.now().strftime('%d %B %Y'))),
        ('Client',           details.get('client', client_name)),
        ('Target',           target_label),
        ('Classification',   'CONFIDENTIAL'),
    ]
    # Switch back to white for the table
    p_spacer = doc.add_paragraph()
    _set_para_bg(p_spacer, NAVY_HEX)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(meta_rows):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
        bg = 'F0F2F8' if i % 2 == 0 else 'FFFFFF'
        for cell in row.cells:
            _set_cell_bg(cell, RGBColor(
                int(bg[0:2], 16), int(bg[2:4], 16), int(bg[4:6], 16)
            ))
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            _set_cell_borders(cell, color='CCCCCC')
        # Bold the label cell
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True

    _page_break(doc)


# ---------------------------------------------------------------------------
# Table of Contents (static — Word can update field on open)
# ---------------------------------------------------------------------------

def _build_toc(doc, findings):
    _heading(doc, 'Table of Contents', level=1, colour=CAAVA_NAVY)

    toc_items = [
        ('1', 'Executive Summary', '3'),
        ('1.1', 'Overview', '3'),
        ('1.2', 'Identified Vulnerabilities', '4'),
        ('2', 'Findings', '5'),
    ]
    for f in findings:
        toc_items.append((f['id'], f['title'], '—'))

    toc_items.append(('A', 'Appendix', '—'))
    toc_items.append(('A.1', 'References and Technical Documentation', '—'))

    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'

    for num, title, page in toc_items:
        row = table.add_row()
        row.cells[0].text = num
        row.cells[1].text = title
        row.cells[2].text = page
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            _set_cell_borders(cell, color='DDDDDD', size='2')

    # Hide table borders
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color='FFFFFF', size='0')

    _para_space(doc, after=4)
    p = doc.add_paragraph()
    _add_run(p, 'CONFIDENTIAL', italic=True, size=9, colour=MID_GREY)
    _page_break(doc)


# ---------------------------------------------------------------------------
# Executive Summary section
# ---------------------------------------------------------------------------

def _build_executive_summary(doc, report_data, findings, chart_png):
    _heading(doc, '1  Executive Summary', level=1, colour=CAAVA_NAVY)
    _heading(doc, '1.1  Overview', level=2, colour=CAAVA_NAVY)

    exec_sum = report_data.get('executive_summary', {})
    overview_text = exec_sum.get('overview', '')
    if overview_text:
        p = doc.add_paragraph(overview_text)
        p.paragraph_format.space_after = Pt(6)

    risk_level = exec_sum.get('risk_level', 'HIGH')
    p_risk = doc.add_paragraph()
    _add_run(p_risk, f'Risk Level: {risk_level}', bold=True, size=11, colour=SEV_COLOURS.get(risk_level.lower(), CAAVA_RED))

    _para_space(doc, after=4)

    # Assessment Details table
    details = exec_sum.get('assessment_details', {})
    if details:
        _heading(doc, 'Assessment Details', level=3, colour=CAAVA_NAVY)
        detail_rows = [
            ('Client',          details.get('client', '')),
            ('Target',          details.get('target', '')),
            ('Assessment Date', details.get('assessment_date', '')),
            ('Assessment Type', details.get('assessment_type', '')),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Item'
        hdr[1].text = 'Details'
        for cell in hdr:
            _set_cell_bg(cell, CAAVA_NAVY)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)

        for label, value in detail_rows:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                _set_cell_borders(cell, color='CCCCCC')

        in_scope = details.get('in_scope', [])
        if in_scope:
            _heading(doc, 'In-Scope Components', level=3, colour=CAAVA_NAVY)
            for item in (in_scope if isinstance(in_scope, list) else [in_scope]):
                p = doc.add_paragraph(style='List Bullet')
                _add_run(p, str(item), size=10)

        out_scope = details.get('out_of_scope', [])
        if out_scope:
            _heading(doc, 'Out-of-Scope', level=3, colour=CAAVA_NAVY)
            for item in (out_scope if isinstance(out_scope, list) else [out_scope]):
                p = doc.add_paragraph(style='List Bullet')
                _add_run(p, str(item), size=10)

    _page_break(doc)


# ---------------------------------------------------------------------------
# Identified Vulnerabilities summary page
# ---------------------------------------------------------------------------

def _build_vuln_summary(doc, findings, chart_png):
    _heading(doc, '1.2  Identified Vulnerabilities', level=2, colour=CAAVA_NAVY)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['#', 'CVSS', 'Description', 'Severity']
    hdr_row = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_row[i].text = h
        _set_cell_bg(hdr_row[i], CAAVA_NAVY)
        for para in hdr_row[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    for f in findings:
        sev = f.get('severity', 'info').lower()
        colour = SEV_COLOURS.get(sev, MID_GREY)
        row = table.add_row()
        row.cells[0].text = f.get('id', '')
        row.cells[1].text = str(f.get('cvss_score', ''))
        row.cells[2].text = f.get('title', '')
        row.cells[3].text = sev.capitalize()
        _set_cell_bg(row.cells[0], colour)
        _set_cell_bg(row.cells[3], colour)
        for ci in [0, 3]:
            for para in row.cells[ci].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
        for ci in [1, 2]:
            for para in row.cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            _set_cell_borders(row.cells[ci], color='CCCCCC')

    _para_space(doc, after=8)

    # Vulnerability Overview heading + chart
    _heading(doc, 'Vulnerability Overview', level=3, colour=CAAVA_NAVY)

    counts = {s: 0 for s in SEV_ORDER}
    for f in findings:
        sev = f.get('severity', 'info').lower()
        if sev in counts:
            counts[sev] += 1

    sev_counts_str = '  '.join(
        f"{s.capitalize()}: {counts[s]}" for s in SEV_ORDER if counts[s] > 0
    )
    p = doc.add_paragraph(f"In the course of this assessment the following vulnerabilities were identified: {sev_counts_str}")
    p.paragraph_format.space_after = Pt(6)

    if chart_png:
        p_chart = doc.add_paragraph()
        p_chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_chart.add_run()
        run.add_picture(io.BytesIO(chart_png), width=Inches(5.5))

    p_fig = doc.add_paragraph('Figure 1 - Distribution of identified vulnerabilities')
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_fig.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GREY

    _page_break(doc)


# ---------------------------------------------------------------------------
# Individual finding section
# ---------------------------------------------------------------------------

def _build_finding(doc, finding):
    fid     = finding.get('id', '')
    title   = finding.get('title', '')
    sev     = finding.get('severity', 'info').lower()
    colour  = SEV_COLOURS.get(sev, MID_GREY)

    # Finding title box (coloured header table)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    _set_cell_bg(table.rows[0].cells[0], colour)
    _set_cell_bg(table.rows[0].cells[1], colour)
    # Merge title row
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    p_title = table.rows[0].cells[0].paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_title, f'{fid}: {title}', bold=True, size=13, colour=WHITE)

    labels_vals = [
        ('Score',         f"{finding.get('cvss_score', '')} ({sev.capitalize()})"),
        ('Vector string', finding.get('cvss_vector', '—')),
        ('Target',        ', '.join(finding.get('affected_components', [])[:1]) or '—'),
        ('References',    ', '.join(finding.get('references', [])[:2]) or '—'),
    ]
    for i, (label, value) in enumerate(labels_vals):
        # Score cell gets severity background
        if i == 0:
            row = table.rows[1]
            _set_cell_bg(row.cells[1], colour)
            row.cells[0].text = label
            p_val = row.cells[1].paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_val, value, bold=True, colour=WHITE, size=10)
        else:
            row = table.add_row() if i > 1 else table.rows[i + 1]
            if i <= 3:
                row = table.rows[i + 1] if i + 1 < len(table.rows) else table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                _set_cell_borders(cell, color='CCCCCC')

    _para_space(doc, after=4)

    # Overview
    overview = finding.get('overview', '')
    if overview:
        _heading(doc, 'Overview', level=3, colour=CAAVA_NAVY)
        p = doc.add_paragraph(overview)
        p.paragraph_format.space_after = Pt(6)

    # Affected Components table
    components = finding.get('affected_components', [])
    if components:
        _heading(doc, 'Affected Components', level=3, colour=CAAVA_NAVY)
        comp_table = doc.add_table(rows=1, cols=2)
        comp_table.style = 'Table Grid'
        comp_hdr = comp_table.rows[0].cells
        comp_hdr[0].text = 'Component'
        comp_hdr[1].text = 'Details'
        for cell in comp_hdr:
            _set_cell_bg(cell, CAAVA_NAVY)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)

        if components and isinstance(components[0], dict):
            for comp in components:
                row = comp_table.add_row()
                row.cells[0].text = comp.get('component', '')
                row.cells[1].text = comp.get('details', '')
        else:
            for comp in components:
                row = comp_table.add_row()
                row.cells[0].text = str(comp)
                row.cells[1].text = ''
        _para_space(doc, after=4)

    # Details
    _heading(doc, 'Details', level=3, colour=CAAVA_NAVY)
    vuln_details = finding.get('vulnerability_details', '')
    if vuln_details:
        _heading(doc, 'Vulnerability Details', level=4, colour=DARK_GREY)
        p = doc.add_paragraph(vuln_details)
        p.paragraph_format.space_after = Pt(6)

    # PoC
    poc = finding.get('proof_of_concept', '')
    if poc:
        _heading(doc, 'Proof of Concept', level=4, colour=DARK_GREY)
        p = doc.add_paragraph(poc)
        p.paragraph_format.space_after = Pt(6)

    # Recommendations
    recs = finding.get('recommendations', {})
    if recs:
        _heading(doc, 'Recommendation', level=3, colour=CAAVA_NAVY)
        for tier_label, tier_key in [('Immediate', 'immediate'), ('Short Term', 'short_term'), ('Long Term', 'long_term')]:
            items = recs.get(tier_key, [])
            if items:
                _heading(doc, tier_label, level=4, colour=DARK_GREY)
                for i, item in enumerate(items, 1):
                    p = doc.add_paragraph(style='List Number')
                    _add_run(p, str(item), size=10)

    _page_break(doc)


# ---------------------------------------------------------------------------
# Appendix
# ---------------------------------------------------------------------------

def _build_appendix(doc, report_data):
    _heading(doc, 'A  Appendix', level=1, colour=CAAVA_NAVY)
    _heading(doc, 'A.1  References and Technical Documentation', level=2, colour=CAAVA_NAVY)

    appendix = report_data.get('appendix', {})

    cve_refs = appendix.get('cve_references', [])
    if cve_refs:
        _heading(doc, 'CVE References', level=3, colour=CAAVA_NAVY)
        for cve in cve_refs:
            _heading(doc, cve.get('id', ''), level=4, colour=CAAVA_RED)
            p_desc = doc.add_paragraph()
            _add_run(p_desc, cve.get('title', ''), bold=True, size=10)
            items = [
                ('CVSS Score', cve.get('cvss', '')),
                ('Description', cve.get('description', '')),
                ('Affected Versions', cve.get('affected_versions', '')),
            ]
            for label, val in items:
                if val:
                    p = doc.add_paragraph(style='List Bullet')
                    _add_run(p, f'{label}: ', bold=True, size=10)
                    _add_run(p, str(val), size=10)

    standards = appendix.get('standards', [])
    if standards:
        _heading(doc, 'Standards and Frameworks', level=3, colour=CAAVA_NAVY)
        for std in standards:
            p = doc.add_paragraph(style='List Bullet')
            _add_run(p, str(std), size=10)

    regulatory = appendix.get('regulatory', [])
    if regulatory:
        _heading(doc, 'Regulatory and Compliance References', level=3, colour=CAAVA_NAVY)
        for reg in regulatory:
            p = doc.add_paragraph(style='List Bullet')
            _add_run(p, str(reg), size=10)

    # Standard disclaimer
    _heading(doc, 'Disclaimer', level=3, colour=CAAVA_NAVY)
    disclaimer = appendix.get('disclaimer',
        'This assessment was conducted with proper authorisation from the client. '
        'All findings are confidential and should be treated as such. '
        'Unauthorised access to the tested systems is illegal. '
        'The vulnerabilities documented in this report should be remediated immediately '
        'to prevent exploitation by malicious actors.')
    p = doc.add_paragraph(disclaimer)
    p.paragraph_format.space_after = Pt(6)

    _para_space(doc, after=4)
    p_contact = doc.add_paragraph()
    _add_run(p_contact, 'Contact and Escalation', bold=True, size=11, colour=CAAVA_NAVY)
    p_contact2 = doc.add_paragraph()
    assessor = os.getenv('ASSESSOR_COMPANY', 'Caava Group')
    _add_run(p_contact2, f'Contact: {assessor} Cyber Security Team\n', size=10)
    _add_run(p_contact2, 'Severity Level: CRITICAL — Immediate executive notification required', size=10)

    _para_space(doc, after=4)
    p_conf = doc.add_paragraph()
    _add_run(p_conf, 'CONFIDENTIAL', italic=True, size=9, colour=MID_GREY)


# ---------------------------------------------------------------------------
# Inner-page header (CAAVA GROUP left | CONFIDENTIAL right)
# ---------------------------------------------------------------------------

def _add_header(doc, target_label):
    """Add branded header to all sections: CAAVA GROUP left, CONFIDENTIAL box right."""
    for section in doc.sections:
        section.different_first_page_header_footer = True   # cover page has no header
        header = section.header
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
        p.clear()

        # Tab stop: left = company name, right = CONFIDENTIAL
        p.paragraph_format.space_after = Pt(0)
        _add_bottom_border(p, 'CCCCCC', '4')

        # Left: CAAVA GROUP
        r_left = p.add_run('CAAVA GROUP')
        r_left.font.color.rgb = CAAVA_NAVY
        r_left.font.bold      = True
        r_left.font.size      = Pt(9)
        r_left.font.name      = 'Calibri'

        # Tab to right
        p.add_run('\t')

        # Right: CONFIDENTIAL in red
        r_right = p.add_run('CONFIDENTIAL')
        r_right.font.color.rgb = CAAVA_RED
        r_right.font.bold      = True
        r_right.font.size      = Pt(9)
        r_right.font.name      = 'Calibri'

        # Set tab stop at right margin
        from docx.oxml import OxmlElement as _OE
        pPr = p._p.get_or_add_pPr()
        tabs = _OE('w:tabs')
        tab  = _OE('w:tab')
        tab.set(qn('w:val'),   'right')
        tab.set(qn('w:pos'),   '9072')   # ~16cm in twips
        tabs.append(tab)
        pPr.append(tabs)


# ---------------------------------------------------------------------------
# Footer on every page (doc title centre | page right | copyright)
# ---------------------------------------------------------------------------

def _add_footer(doc, target_label):
    """Footer: doc title center, copyright right."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        _add_top_border(p, 'CCCCCC', '4')

        assessor = os.getenv('ASSESSOR_COMPANY', 'Caava Group')
        year     = datetime.now().year

        r_title = p.add_run(f'{target_label}    ')
        r_title.font.italic = True
        r_title.font.size   = Pt(8)
        r_title.font.color.rgb = MID_GREY
        r_title.font.name   = 'Calibri'

        p.add_run('\t')

        r_copy = p.add_run(f'© {year} {assessor}. All Rights Reserved.')
        r_copy.font.size      = Pt(8)
        r_copy.font.color.rgb = MID_GREY
        r_copy.font.name      = 'Calibri'

        # Right-align tab
        from docx.oxml import OxmlElement as _OE
        pPr = p._p.get_or_add_pPr()
        tabs = _OE('w:tabs')
        tab  = _OE('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab)
        pPr.append(tabs)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def build_docx(report_data: dict, findings: list, target_url: str,
               client_name: str, output_path: Path) -> Path:
    """
    report_data : parsed Claude JSON
    findings    : list of finding dicts (same as report_data['findings'])
    target_url  : scanned URL
    client_name : e.g. "Caava Group"
    output_path : where to save the .docx
    """
    doc = Document()

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    target_label = report_data.get('executive_summary', {}).get(
        'assessment_details', {}).get('target', target_url)

    chart_png = _build_severity_chart(findings)

    _build_cover(doc, report_data, target_url, client_name)
    _build_toc(doc, findings)
    _build_executive_summary(doc, report_data, findings, chart_png)
    _build_vuln_summary(doc, findings, chart_png)

    # Findings section heading
    _heading(doc, '2  Findings', level=1, colour=CAAVA_NAVY)
    _para_space(doc)
    for finding in findings:
        _build_finding(doc, finding)

    _build_appendix(doc, report_data)
    _add_header(doc, target_label)
    _add_footer(doc, target_label)

    doc.save(str(output_path))
    return output_path
